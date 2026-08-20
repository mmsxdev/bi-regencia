# -*- coding: utf-8 -*-
"""
Gera o documento de apoio de 1 pagina (handout) em Word.

Uso:
    python gerar_resumo.py

Requisitos: python-docx
Saida: RESUMO_1_PAGINA.docx (raiz do projeto)
"""
import os
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

RED = RGBColor(0xD7, 0x19, 0x20)
BLUE = RGBColor(0x2E, 0x90, 0xFA)
DARK = RGBColor(0x0B, 0x0F, 0x14)
GRAY = RGBColor(0x6B, 0x72, 0x80)

OUT = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "RESUMO_1_PAGINA.docx")


def fill(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.2)
    sec.top_margin = Cm(1.6)
    sec.bottom_margin = Cm(1.6)

    def para(text="", size=10.5, bold=False, color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_after=4, italic=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
        r.font.name = "Segoe UI"
        return p

    # Cabeçalho
    para("Complexo de Educação, Tecnologia, Inovação e Saúde Paulo Vargas", 9, True, GRAY)
    para("BI DE REGÊNCIA — SÍNTESE EXECUTIVA", 21, True, RED, WD_ALIGN_PARAGRAPH.LEFT, 2)
    para("Regência dos instrutores em sala de aula · Complexo de Educação, Tecnologia, Inovação e Saúde Paulo Vargas · Quadro 2026",
         11, False, DARK, WD_ALIGN_PARAGRAPH.LEFT, 10)

    # Ficha de indicadores
    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    dados = [
        ("INSTRUTORES", "45", "no quadro (total desde ago/26)"),
        ("REGÊNCIA MÉDIA", "61%", "média das regências mensais"),
        ("HORAS-AULA", "37,3 mil", "soma dos 12 meses (2026)"),
        ("ABAIXO DE 50%", "12 instr.", "27% do quadro em alerta"),
    ]
    for j, (lab, val, hint) in enumerate(dados):
        c1 = table.cell(0, j)
        c2 = table.cell(1, j)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(lab)
        r1.font.size = Pt(8.5)
        r1.font.bold = True
        r1.font.color.rgb = GRAY
        r1.font.name = "Segoe UI"
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(val)
        r2.font.size = Pt(20)
        r2.font.bold = True
        r2.font.color.rgb = DARK
        r2.font.name = "Segoe UI"
    para("", size=4, space_after=6)

    def heading(t, color=BLUE):
        para(t, 12.5, True, color, WD_ALIGN_PARAGRAPH.LEFT, 3)

    def bullet(t, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.4)
        r0 = p.add_run("•  ")
        r0.font.size = Pt(9.5)
        r0.font.bold = True
        r0.font.color.rgb = RED
        r0.font.name = "Segoe UI"
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.size = Pt(9.8)
            rb.font.bold = True
            rb.font.color.rgb = DARK
            rb.font.name = "Segoe UI"
        r = p.add_run(t)
        r.font.size = Pt(9.8)
        r.font.color.rgb = DARK
        r.font.name = "Segoe UI"

    heading("O QUE É", RED)
    para("Painel interativo (Streamlit) que consolida as horas-aula dos instrutores do quadro a partir "
         "da aba CONSOLIDADO da planilha de regência, com atualização automática.", 10, space_after=6)

    heading("PRINCIPAIS ACHADOS (2026)")
    bullet("soma dos meses (37.306 h) difere da coluna ANO (36.458 h) — padronizar cálculo na planilha.",
           "Dois totais: ")
    bullet("12 instrutores (27%) abaixo de 50% de regência; 2 acima de 100% (referência).", "Alerta: ")
    bullet("área unificada (13 instrutores), 7 abaixo de 50% — prioridade com a coordenação.",
           "Manutenção automotiva: ")
    bullet("julho (18% — férias), janeiro (41%) e dezembro (37%); base produtiva: mar–set (78–82%).",
           "Sazonalidade: ")
    bullet("Gráfica editorial (74,7%) e Alimentos e bebidas (73,0%); ROMULO FLORIANO 109,5%.",
           "Referências boas: ")
    bullet("erros de digitação e rótulos antigos são unificados pelo próprio BI (ex.: \"(JD)\" → "
           "Manutenção automotiva); falta preencher POLO/LOCAL na planilha para todos.", "Qualidade dos dados: ")

    heading("PLANO DE AÇÃO")
    bullet("validar casos de baixa regência (por que, o que fazer) — prioridade Manutenção automotiva.",
           "1) ")
    bullet("padronizar planilha: preencher POLO/LOCAL (já suportado pelo BI), totais anuais e "
           "lançamentos mensais.", "2) ")
    bullet("definir metas de regência por período, respeitando férias e sazonalidade.", "3) ")
    bullet("acompanhamento mensal no painel, com reavaliação do mesmo indicador.", "4) ")

    para("", size=4, space_after=4)
    para(f"Fonte: planilha REGÊNCIA - INSTRUTORES DO QUADRO 2026.xlsx (aba CONSOLIDADO) processada pelo "
         f"BI de Regência. Gerado em {date.today().strftime('%d/%m/%Y')} — os números podem ser atualizados "
         "pela coordenação na planilha-fonte.", 8.5, False, GRAY, WD_ALIGN_PARAGRAPH.LEFT, 0)


def main():
    doc = Document()
    fill(doc)
    doc.save(OUT)
    print("OK", OUT)


if __name__ == "__main__":
    main()